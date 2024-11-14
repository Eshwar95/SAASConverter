#!/usr/bin/env python3


"""
Created on Wed Oct 30 16:20:01 2024

@author: ayshwarya
"""
import io
from io import BytesIO
import streamlit as st
import openai
import pandas as pd
from docx import Document
import os
import tempfile
import getpass
#from CodeConversion import extract_zip
from pathlib import Path
import rag_code as rag

import zipfile
# Set your OpenAI API key here
deployment_name="gpt-35-turbo"
openai.api_type = "azure"
openai.api_key = "" #os.getenv("AZURE_OPENAI_API_KEY")  # Alternatively, paste your key directly: "your_api_key"
openai.api_base = ""  # Replace with your Azure OpenAI endpoint
openai.api_version = "2024-07-01-preview"  # Ensure the API version matches the one in your Azure Portal



def read_file(file_path):
	"""Read data from a file."""
	try:
		# Detect file extension and read accordingly
		if file_path.endswith('.txt') or file_path.endswith('.js') or file_path.endswith('.py'):
			with open(file_path, 'r', encoding='utf-8') as file:
				content = file.read()
		elif file_path.endswith('.json'):
			import json
			with open(file_path, 'r', encoding='utf-8') as file:
				content = json.load(file)
		elif file_path.endswith('.csv'):
			import csv
			with open(file_path, 'r', encoding='utf-8') as file:
				reader = csv.reader(file)
				content = list(reader)
		elif file_path.endswith('.xml'):
			import xml.etree.ElementTree as ET
			tree = ET.parse(file_path)
			content = tree.getroot()
		elif file_path.endswith('.yaml') or file_path.endswith('.yml'):
			import yaml
			with open(file_path, 'r', encoding='utf-8') as file:
				content = yaml.safe_load(file)
		else:
			raise ValueError("Unsupported file extension.")
		return content
	except Exception as e:
		print(f"Error reading file: {e}")
		return None
	
def generate_summary(data):
	"""Generate sentences using OpenAI's GPT model."""
	prompt = f"Describe what the python code does in 2000 words and seperate it by using headings:\n\n{data}\n\nSummary:"
	
	response = openai.ChatCompletion.create(
		engine="gpt-35-turbo",  # You can change the model if needed
		messages=[{"role": "user", "content": prompt}]
	)  
	
	return response.choices[0].message.content


def create_word_document(content, output_file):
	"""Create a Word document with the generated content."""
	doc = Document()
	doc.add_heading('Generated Summary', level=1)
	doc.add_paragraph(content)
	doc.save(output_file)
	
	# Save the document to a BytesIO object
	doc_io = BytesIO()
	doc.save(doc_io)
	doc_io.seek(0)  # Reset the pointer to the beginning of the file
	
	return doc_io


def main(file, file_name, output_word_file):
	"""Main function to read Excel, generate summary, and create Word document."""
	# Step 1: Read data from Excel
	data = read_file(file)
	
	# Step 2: Generate summary using OpenAI API
	summary = generate_summary(file)
	# print(file)
	
	# Step 3: Create a Word document with the summary
	doc = create_word_document(summary, output_word_file)
	return doc


def app():
    tab1, tab2 = st.tabs(["Summary Generation", "Retrieval-Augmented Generation"])
    with tab1:
        st.write('# 📄 Documentation Generator')
        st.write("Upload a ZIP file containing code files to automatically generate documentation summaries.")
        
        uploaded_file = st.file_uploader(label='📂 Upload your zip file here', type=['zip'])
        if uploaded_file:
            # Progress bar to show upload progress
            with st.spinner("Extracting files and preparing for documentation..."):
                with tempfile.NamedTemporaryFile(delete=False) as tmp_zip:
                    tmp_zip.write(uploaded_file.read())
                    zip_file_path = tmp_zip.name
                    
                # Extract files and read contents
                extracted_files = read_zip_contents(zip_file_path)
                
            # Initialize session state for storing generated files if not already present
            if 'generated_files' not in st.session_state:
                st.session_state.generated_files = {}
                
            if st.button('🚀 Generate Documentation'):
                with st.spinner("Generating summaries..."):
                    # Generate documents for each extracted file and store in session state
                    for file_name, contents in extracted_files.items():
                        # Call `main()` to generate the summary for each file and get a BytesIO document
                            doc = main(contents, file_name, 'document.docx')
                            st.session_state.generated_files[file_name] = doc
                            
                st.success("✅ Documentation generated successfully! Expand below to download your summaries.")
                
        # Use an expander to contain all download buttons neatly
        if 'generated_files' in st.session_state and st.session_state.generated_files:
            with st.expander("📥 Download Summaries"):
                for file_name, doc in st.session_state.generated_files.items():
                    st.write(f"#### {file_name} Summary")
                    st.download_button(
                        label=f"Download {file_name} Summary",
                        data=doc,
                        file_name=f"summary_{file_name}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_{file_name}"
                    )
                    st.write("")  # Adds spacing between buttons
					
    with tab2:
        print(1)
        #rag.app()

			
def generate_documents(extracted_files):
	# Temporary directory to store individual .docx files
	temp_dir = 'temp_docs'
	os.makedirs(temp_dir, exist_ok=True)
	
	# Create a zip in memory
	zip_buffer = io.BytesIO()
	
	# Create a ZipFile object in memory
	with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
		for file_name, contents in extracted_files.items():
			# Generate the document (assuming `main` returns a .docx file's content as a BytesIO object)
			doc_io = main(contents, file_name, 'document.docx')
			
			# Retrieve bytes from the BytesIO object
			doc_bytes = doc_io.getvalue()
			
			# Temporary file name for each document in the zip
			doc_filename = f"summary_{file_name}.docx"
			
			# Write the document content to the zip file
			zip_file.writestr(doc_filename, doc_bytes)
			
	# Seek to the beginning of the in-memory zip file
	zip_buffer.seek(0)
	
	return zip_buffer



def read_zip_contents(zip_file_path):
	"""Extracts the contents of the uploaded ZIP file and reads each code file."""
	extracted_files = {}
	
	with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
		# Extract all files in the ZIP to a temporary directory
		temp_dir = tempfile.mkdtemp()
		zip_ref.extractall(temp_dir)
		
		# Read the contents of each file in the ZIP
		for file_name in zip_ref.namelist():
			file_path = os.path.join(temp_dir, file_name)
			
			file_contents = read_file(file_path)
			# print(file_contents)
			extracted_files[file_name] = file_contents
			
	return extracted_files

# Run the app
if __name__ == "__main__":
    app()
