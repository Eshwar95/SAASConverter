import os
import zipfile
import openai
import time
import tempfile
import pandas as pd
from io import BytesIO
import streamlit as st
from pathlib import Path
from docx import Document
  # Ensure python-docx is installed

# Set up OpenAI API configuration
openai.api_key = ""  # Securely set this

# Streamlit custom styles
st.set_page_config(page_title="File Converter & Unit Test Generator", layout="centered")
st.markdown("""
    <style>
    .stApp { background-color: #f4f4f9; color: #333; }
    .main-title { font-size: 36px; font-weight: 600; color: #336699; margin-top: 10px; margin-bottom: 20px; }
    .sub-title { font-size: 20px; font-weight: 400; color: #555; }
    .download-btn { background-color: #4CAF50; color: white; font-weight: 600; }
    </style>
""", unsafe_allow_html=True)

# Sidebar menu for navigation
st.sidebar.title("Menu")
menu_option = st.sidebar.radio("Choose a section:", ["Conversion", "Document Creation"])

# Function to extract zip files
def extract_zip(zip_path, extract_to):
    with zipfile.ZipFile(zip_path, 'r') as zip_ref:
        zip_ref.extractall(extract_to)

# Function to read and convert Excel content to code format
def read_excel(file_path):
    try:
        excel_data = pd.read_excel(file_path)
        return excel_data.to_string()
    except Exception as e:
        st.warning(f"Error reading Excel file: {e}")
        return None

# Read file content with encoding handling
def read_file(file_path, file_type):
    if file_type == 'xlsx':
        return read_excel(file_path)
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding='ISO-8859-1') as file:
            return file.read()

# Convert code using OpenAI API
def get_code_from_openai(content, language):
    prompt = f"Convert the following code to {language} and provide only the code without any explanations and comments:\n\n{content}"
    try:
        messages = [
            {"role": "system", "content": f"You are an AI that generates {language} code."},
            {"role": "user", "content": prompt}
        ]
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=messages,
            max_tokens=800,
            temperature=0.7
        )
        return response.choices[0].message['content'].strip()
    except Exception as e:
        st.warning(f"Error with OpenAI API: {e}")
        return None

# Generate unit tests
def generate_unit_test(filecontent, language):
    try:
        prompt = f"Generate unit tests for the following {language} code. Provide only the code without any explanations:\n\n{filecontent}"
        messages = [
            {"role": "system", "content": f"You are an AI that generates {language} unit tests."},
            {"role": "user", "content": prompt}
        ]
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=messages,
            max_tokens=800,
            temperature=0.7
        )
        return response.choices[0].message['content'].strip()
    except Exception as e:
        st.warning(f"Error generating unit tests: {e}")
        return None

# Process files in zip and retrieve converted code
def process_files(zip_file, language, generate_tests):
    with tempfile.TemporaryDirectory() as temp_dir:
        extract_zip(zip_file, temp_dir)
        converted_files = {}
        supported_extensions = ['sas', 'xlsx', 'csv', 'txt', 'py', 'js', 'cs', 'java']  # Add more if needed
        files = [f for f in os.listdir(temp_dir) if any(f.endswith(ext) for ext in supported_extensions)]
        total_steps = len(files) * (2 if generate_tests else 1)
        progress = st.progress(0)

        for i, file_name in enumerate(files):
            file_path = os.path.join(temp_dir, file_name)
            file_type = Path(file_name).suffix[1:]
            content = read_file(file_path, file_type)

            # Code conversion step
            if content:
                code = get_code_from_openai(content, language)
                if code:
                    converted_files[file_name] = code
                    time.sleep(1)  # Delay for rate limits

            # Unit test generation step
            if generate_tests and code:
                unit_test_code = generate_unit_test(code, language)
                if unit_test_code:
                    converted_files[file_name + "_test"] = unit_test_code
                progress.progress((i + 1) * 2 / total_steps)  # Update for unit test step

            # Update progress
            progress.progress((i + 1) / total_steps)

        progress.empty()
        return converted_files

# Document Creation Section Functions
# Function to read file content for document creation
def read_file_for_document(file_path, file_type):
    if file_type == 'xlsx':
        return read_excel(file_path)
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding='ISO-8859-1') as file:
            return file.read()

# Process files in zip for document creation
def process_files_for_document(zip_file):
    with tempfile.TemporaryDirectory() as temp_dir:
        extract_zip(zip_file, temp_dir)
        all_content = []
        supported_extensions = ['sas', 'xlsx', 'csv', 'txt', 'py', 'js', 'cs', 'java']  # Add more if needed
        files = [f for f in os.listdir(temp_dir) if any(f.endswith(ext) for ext in supported_extensions)]

        for file_name in files:
            file_path = os.path.join(temp_dir, file_name)
            file_type = Path(file_name).suffix[1:]
            content = read_file_for_document(file_path, file_type)
            if content:
                all_content.append(content)

        return "\n\n".join(all_content)

# Main Application Logic
if menu_option == "Conversion":
    # User upload zip file for conversion
    st.header("Code converter and  test case generator")
    uploaded_file = st.file_uploader("Upload a zip file with various code files", type="zip", key="converter_uploader")
    language = st.selectbox("Choose the target language for conversion:", ["Python", "C#", "Java", "JavaScript"])
    generate_tests = st.checkbox("Generate unit tests for the chosen language")

    # Determine the appropriate file extension based on the language
    file_extension_map = {"Python": "py", "C#": "cs", "Java": "java", "JavaScript": "js"}
    file_extension = file_extension_map[language]

    # Conversion and unit test generation
    if uploaded_file and language:
        if st.button("Start Conversion"):
            st.info("Processing files...")
            with tempfile.NamedTemporaryFile(delete=False) as tmp_zip:
                tmp_zip.write(uploaded_file.read())
                zip_file_path = tmp_zip.name
            converted_files = process_files(zip_file_path, language, generate_tests)

            if converted_files:
                output_zip = BytesIO()
                with zipfile.ZipFile(output_zip, 'w') as zf:
                    for file_name, code in converted_files.items():
                        if "_test" not in file_name:
                            new_file_name = os.path.splitext(file_name)[0] + f"_converted.{file_extension}"
                            zf.writestr(new_file_name, code)
                            with st.expander(f"Converted Code: {file_name} ({language})"):
                                st.code(code, language=language.lower())

                        elif generate_tests:
                            new_test_name = os.path.splitext(file_name)[0] + f"_test.{file_extension}"
                            zf.writestr(new_test_name, code)

                    output_zip.seek(0)

                st.download_button(
                    label="📥 Download Converted Files and Unit Tests",
                    data=output_zip,
                    file_name="converted_and_tests.zip",
                    mime="application/zip",
                    key="download-btn"
                )
                st.success("Conversion complete! Download your files, including unit tests if selected.")

elif menu_option == "Document Creation":
    st.subheader("Create Document")
    # User upload zip file for document creation
    uploaded_doc_file = st.file_uploader("Upload a zip file with various code files for document creation", type="zip", key="doc_uploader")
    
    # Document summary generation
    if uploaded_doc_file and st.button("Generate Summary Document"):
        st.info("Generating summary...")
        with tempfile.NamedTemporaryFile(delete=False) as tmp_zip:
            tmp_zip.write(uploaded_doc_file.read())
            zip_file_path = tmp_zip.name

        combined_content = process_files_for_document(zip_file_path)

        if combined_content:
            # Generate summary using OpenAI
            summary_prompt = f"Summarize what code does and minimun words for summary are 2000 and use RAG implementation:\n\n{combined_content}"
            try:
                messages = [
                    {"role": "system", "content": "You are an AI that generates summaries."},
                    {"role": "user", "content": summary_prompt}
                ]
                response = openai.ChatCompletion.create(
                    model="gpt-4",
                    messages=messages,
                    max_tokens=800,
                    temperature=0.7
                )
                summary = response.choices[0].message['content'].strip()
                
                # Create a Word document
                doc = Document()
                doc.add_heading('Combined Summary', level=1)
                doc.add_paragraph(summary)

                # Save document to a temporary file
                temp_doc_path = tempfile.mktemp(suffix='.docx')
                doc.save(temp_doc_path)

                # Provide download button
                with open(temp_doc_path, "rb") as f:
                    st.download_button(
                        label="📥 Download Summary Document",
                        data=f,
                        file_name="summary_document.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                st.success("Summary generated! You can download the document.")

            except Exception as e:
                st.warning(f"Error generating summary: {e}")
        else:
            st.warning("No content found to summarize.")
