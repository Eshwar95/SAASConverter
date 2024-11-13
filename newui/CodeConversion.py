import os
import zipfile
import openai
import tempfile
import pandas as pd
from io import BytesIO
import streamlit as st
from pathlib import Path
from docx import Document

# Set your OpenAI API key here
openai.api_key = "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"

# Function to extract zip files
def extract_zip(zip_path, extract_to):
    with zipfile.ZipFile(zip_path, 'r') as zip_ref:
        zip_ref.extractall(extract_to)

# Function to read and convert Excel content to code format
def read_excel(file_path):
    try:
        excel_data = pd.read_excel(file_path)
        return excel_data.to_string()
    except Exception as e:
        st.warning(f"Error reading Excel file: {e}")
        return None

# Read file content with encoding handling
def read_file(file_path, file_type):
    if file_type == 'xlsx':
        return read_excel(file_path)
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding='ISO-8859-1') as file:
            return file.read()

# Convert code using OpenAI API
def get_code_from_openai(content, language):
    prompt = f"Convert the following code to {language} and provide only the code without any explanations and comments:\n\n{content}"
    try:
        messages = [
            {"role": "system", "content": f"You are an AI that generates {language} code."},
            {"role": "user", "content": prompt}
        ]
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=messages,
            max_tokens=800,
            temperature=0.7
        )
        return response.choices[0].message['content'].strip()
    except Exception as e:
        st.warning(f"Error with OpenAI API: {e}")
        return None

# Generate unit tests
def generate_unit_test(filecontent, language):
    try:
        prompt = f"Generate unit tests for the following {language} code. Provide only the code without any explanations:\n\n{filecontent}"
        messages = [
            {"role": "system", "content": f"You are an AI that generates {language} unit tests."},
            {"role": "user", "content": prompt}
        ]
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=messages,
            max_tokens=800,
            temperature=0.7
        )
        return response.choices[0].message['content'].strip()
    except Exception as e:
        st.warning(f"Error generating unit tests: {e}")
        return None

# Summarize automation
def generate_summary(content):
    try:
        prompt = f"Provide a concise summary of the following content:\n\n{content}"
        messages = [
            {"role": "system", "content": "You are an AI that generates summaries."},
            {"role": "user", "content": prompt}
        ]
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=messages,
            max_tokens=800,
            temperature=0.7
        )
        return response.choices[0].message['content'].strip()
    except Exception as e:
        st.warning(f"Error generating summary: {e}")
        return None

# Create documentation
def generate_documentation(content):
    try:
        prompt = f"Create detailed documentation for the following content:\n\n{content}"
        messages = [
            {"role": "system", "content": "You are an AI that generates documentation."},
            {"role": "user", "content": prompt}
        ]
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=messages,
            max_tokens=800,
            temperature=0.7
        )
        return response.choices[0].message['content'].strip()
    except Exception as e:
        st.warning(f"Error generating documentation: {e}")
        return None

# Process files for document creation
def process_files_for_document(zip_file):
    with tempfile.TemporaryDirectory() as temp_dir:
        extract_zip(zip_file, temp_dir)
        all_content = []
        supported_extensions = ['sas', 'xlsx', 'csv', 'txt', 'py', 'js', 'cs', 'java']
        files = [f for f in os.listdir(temp_dir) if any(f.endswith(ext) for ext in supported_extensions)]

        for file_name in files:
            file_path = os.path.join(temp_dir, file_name)
            file_type = Path(file_name).suffix[1:]
            content = read_file(file_path, file_type)
            if content:
                all_content.append(content)

        return "\n\n".join(all_content)

# Create a Word document
def create_docx(content, title):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# Process files for conversion and test generation
def process_files(zip_file, languages, generate_tests):
    converted_files = {}
    with tempfile.TemporaryDirectory() as temp_dir:
        extract_zip(zip_file, temp_dir)
        supported_extensions = ['sas', 'xlsx', 'csv', 'txt', 'py', 'js', 'cs', 'java']
        files = [f for f in os.listdir(temp_dir) if any(f.endswith(ext) for ext in supported_extensions)]

        for file_name in files:
            file_path = os.path.join(temp_dir, file_name)
            file_type = Path(file_name).suffix[1:]
            content = read_file(file_path, file_type)

            if content:
                for language in languages:
                    code = get_code_from_openai(content, language)
                    if code:
                        converted_files[f"{file_name}_to_{language}"] = code
                    if generate_tests and code:
                        unit_test_code = generate_unit_test(code, language)
                        if unit_test_code:
                            converted_files[f"{file_name}_to_{language}_test"] = unit_test_code

        return converted_files

# Main app function
def app():
    st.write('# Code Conversion and Automation Tool')

    uploaded_file = st.file_uploader(label='Upload a zip file with code files')
    languages = ['Python', 'C#', 'Java', 'JavaScript']
    selected_languages = st.multiselect("Select target language(s) for conversion", languages)
    automations = ['Generate unit tests for selected language(s)', 'Summarize automation', 'Create Documentation']
    selected_automations = st.multiselect("Select automation(s)", automations)

    generate_tests = 'Generate unit tests for selected language(s)' in selected_automations
    summarize_automation = 'Summarize automation' in selected_automations
    create_documentation = 'Create Documentation' in selected_automations

    # Initialize session state for persistence
    if 'converted_files' not in st.session_state:
        st.session_state.converted_files = None
        st.session_state.summary_docx = None
        st.session_state.documentation_docx = None
        st.session_state.expanders_open = {}

    if uploaded_file:
        with tempfile.NamedTemporaryFile(delete=False) as tmp_zip:
            tmp_zip.write(uploaded_file.read())
            zip_file_path = tmp_zip.name

        if st.button('Generate Conversion and Tests'):
            with st.spinner("Processing files..."):
                converted_files = process_files(zip_file_path, selected_languages, generate_tests)
                combined_content = "\n\n".join(converted_files.values())

                if summarize_automation:
                    summary_content = generate_summary(combined_content)
                    st.session_state.summary_docx = create_docx(summary_content, "Summary")
                if create_documentation:
                    documentation_content = generate_documentation(combined_content)
                    st.session_state.documentation_docx = create_docx(documentation_content, "Documentation")

                # Store converted files in session state
                output_zip = BytesIO()
                with zipfile.ZipFile(output_zip, 'w') as zf:
                    for file_name, code in converted_files.items():
                        zf.writestr(f"{file_name}.txt", code)

                        # Track expander state in session state
                        if file_name not in st.session_state.expanders_open:
                            st.session_state.expanders_open[file_name] = False

                output_zip.seek(0)
                st.session_state.converted_files = output_zip

            st.success("Conversion and Test Generation Complete!")

    # Display download buttons if files are ready
    if st.session_state.converted_files:
        st.download_button("📥 Download Converted Files and Unit Tests", data=st.session_state.converted_files, file_name="converted_and_tests.zip")
    if st.session_state.summary_docx:
        st.download_button("📥 Download Summary Document", data=st.session_state.summary_docx, file_name="summary.docx")
    if st.session_state.documentation_docx:
        st.download_button("📥 Download Documentation Document", data=st.session_state.documentation_docx, file_name="documentation.docx")

    # Show expanders with converted code
    if st.session_state.converted_files:
        with zipfile.ZipFile(st.session_state.converted_files, 'r') as zf:
            for file_name in zf.namelist():
                with zf.open(file_name) as file:
                    code_content = file.read().decode('utf-8')
                    if file_name not in st.session_state.expanders_open:
                        st.session_state.expanders_open[file_name] = False
                    with st.expander(f"Converted Code: {file_name}", expanded=st.session_state.expanders_open[file_name]):
                        st.code(code_content)
                        # Toggle expander state
                        st.session_state.expanders_open[file_name] = True

# Run the app
if __name__ == "__main__":
    app()
