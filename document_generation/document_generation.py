#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Wed Oct 30 16:20:01 2024

@author: ayshwarya
"""

import openai
import pandas as pd
from docx import Document
import os
import getpass

deployment_name="gpt-35-turbo-16k"
openai.api_type = "azure"
openai.api_key = "test" #os.getenv("AZURE_OPENAI_API_KEY")  # Alternatively, paste your key directly: "your_api_key"
openai.api_base = "test"  # Replace with your Azure OpenAI endpoint
openai.api_version = "2024-06-01"  # Ensure the API version matches the one in your Azure Portal

if "AZURE_OPENAI_API_KEY" not in os.environ:
   os.environ["AZURE_OPENAI_API_KEY"] = getpass.getpass(
       "test"
   )
os.environ["AZURE_OPENAI_ENDPOINT"] = "test"

def read_file(file_path):
    """Read data from a file."""
    try:
        # Detect file extension and read accordingly
        if file_path.endswith('.txt') or file_path.endswith('.js') or file_path.endswith('.py'):
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
        elif file_path.endswith('.json'):
            import json
            with open(file_path, 'r', encoding='utf-8') as file:
                content = json.load(file)
        elif file_path.endswith('.csv'):
            import csv
            with open(file_path, 'r', encoding='utf-8') as file:
                reader = csv.reader(file)
                content = list(reader)
        elif file_path.endswith('.xml'):
            import xml.etree.ElementTree as ET
            tree = ET.parse(file_path)
            content = tree.getroot()
        elif file_path.endswith('.yaml') or file_path.endswith('.yml'):
            import yaml
            with open(file_path, 'r', encoding='utf-8') as file:
                content = yaml.safe_load(file)
        else:
            raise ValueError("Unsupported file extension.")
        return content
    except Exception as e:
        print(f"Error reading file: {e}")
        return None

def generate_summary(data):
    """Generate sentences using OpenAI's GPT model."""
    prompt = f"Describe what the python code does in 2000 words and seperate it by using headings:\n\n{data}\n\nSummary:"

    response = openai.ChatCompletion.create(
        engine="gpt-35-turbo-16k",  # You can change the model if needed
        messages=[{"role": "user", "content": prompt}]
    )

    return response.choices[0].message.content


def create_word_document(content, output_file):
    """Create a Word document with the generated content."""
    doc = Document()
    doc.add_heading('Generated Summary', level=1)
    doc.add_paragraph(content)
    doc.save(output_file)

def main(file, output_word_file):
    """Main function to read Excel, generate summary, and create Word document."""
    # Step 1: Read data from Excel
    data = read_file(file)

    # Step 2: Generate summary using OpenAI API
    summary = generate_summary(data)

    # Step 3: Create a Word document with the summary
    create_word_document(summary, output_word_file)

if __name__ == "__main__":
    file_path = '/Users/ayshwarya/Downloads/captcha_generator.py'
    output_word_file_path = 'file.docx'

    main(file_path, output_word_file_path)
